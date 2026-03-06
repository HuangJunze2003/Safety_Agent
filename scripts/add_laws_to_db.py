import os
import argparse
import sys
import re
from pathlib import Path
import subprocess
from langchain_core.documents import Document
import fitz
from docx import Document as DocxDocument

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT / "src") not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT / "src"))

from agent.workflow import AgentConfig, LegalClauseRetriever

def extract_text_from_file(file_path: Path) -> str:
    """自动判断文件类型并提取文字"""
    ext = file_path.suffix.lower()
    text = ""
    try:
        if ext == ".txt":
            text = file_path.read_text(encoding="utf-8")
        elif ext == ".pdf":
            with fitz.open(file_path) as doc:
                text = "\n".join(page.get_text() for page in doc)
        elif ext == ".docx":
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif ext == ".doc":
            # 使用 antiword 提取老式二进制的 .doc 格式
            result = subprocess.run(['antiword', str(file_path)], capture_output=True, text=True, check=True)
            text = result.stdout
    except Exception as e:
        print(f"读取文件 {file_path.name} 失败: {e}")
    # 移除首尾空白
    return text.strip()

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> list[str]:
    """
    更稳健的法条切分逻辑。
    """
    # 匹配换行开头的 "第X条"，支持空格
    # 兼容常见的几种格式：第十条、第一章、第一节
    pattern = r'(\n\s*(?:第[一二三四五六七八九十百]+[条章节]))'
    parts = re.split(pattern, text)
    
    if len(parts) <= 1:
        # 如果没有找到明显的层级标志，则常规切分
        res = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            res.append(text[start:end])
            start = end - overlap if end < len(text) else len(text)
        return res

    chunks = []
    # parts[0] 可能是标题等前置内容
    if parts[0].strip():
        # 如果标题部分包含文件名或法律名，赋予更高权重
        chunks.append(parts[0].strip())
        
    i = 1
    recent_chapter = "" # 记录最近的章节信息
    while i < len(parts):
        header = parts[i]
        content = parts[i+1] if i+1 < len(parts) else ""
        
        # 如果是章节，记录下来，并加到后续法条中作为上下文
        if "章" in header or "节" in header:
            recent_chapter = header + content.strip().split('\n')[0] # 取第一行
            chunks.append((header + content).strip())
        else:
            # 这是一个 "条"，把最近的章节信息一起带入
            combined = (recent_chapter + "\n" + header + content).strip()
            chunks.append(combined)
        i += 2
            
    return chunks

def main():
    parser = argparse.ArgumentParser(description="将外部的法律法规文件（TXT格式）导入到法规检索数据库中")
    parser.add_argument("--laws-dir", type=Path, default=PROJECT_ROOT / "data/data_raw/laws", help="存放法律TXT文件的目录")
    args = parser.parse_args()
    
    laws_dir = args.laws_dir
    if not laws_dir.exists():
        laws_dir.mkdir(parents=True, exist_ok=True)
        print(f"📁 已创建目录 {laws_dir}，请将法律法规文件放入该目录中再运行此脚本。")
        return
        
    law_files = []
    # 使用 rglob 递归搜索该目录以及所有子目录下的指定格式文件
    for ext in ["*.txt", "*.TXT", "*.pdf", "*.PDF", "*.docx", "*.DOCX", "*.doc", "*.DOC"]:
        law_files.extend(list(laws_dir.rglob(ext)))
        
    if not law_files:
        print(f"⚠️ 在 {laws_dir} 中未找到任何 TXT/PDF/Word 文件。请放入经过整理的法律文本。")
        return
        
    print("⏳ 正在初始化法条检索器与向量数据库...")
    config = AgentConfig(project_root=str(PROJECT_ROOT))
    retriever = LegalClauseRetriever(
        project_root=config.project_root,
        metadata_path=config.cases_metadata_path,
        persist_directory=config.legal_db_dir,
        embedding_model_name=config.bge_model_name # 使用 config 中定义的默认模型
    )
    
    # 确保基础数据已入库
    retriever._ensure_indexed()
    
    if not retriever.store:
        print("❌ 离线模式下无法获取 embeddings 模型，无法入库。请关闭离线模式环境变量。")
        return
    
    new_docs = []
    print(f"🚀 开始处理 {len(law_files)} 个法律文件...")
    for file_path in law_files:
        print(f"  📖 正在读取并切分: {file_path.name}")
        content = extract_text_from_file(file_path)
        
        if not content.strip():
            print(f"  ⚠️ 文件 {file_path.name} 未能提取到有效文本或内容为空。")
            continue
            
        chunks = chunk_text(content, chunk_size=400, overlap=50)
        
        for i, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            # 在内容开头注入法律名称和章节信息辅助检索 (Self-Correction: 增强命中率)
            tagged_content = f"《{file_path.stem}》\n{chunk}"
            doc = Document(
                page_content=tagged_content,
                metadata={
                    "source_file": file_path.name,
                    "chunk_id": i,
                    "type": "external_law"
                }
            )
            new_docs.append(doc)
            
    if new_docs:
        print(f"💾 正在将 {len(new_docs)} 个文本片段向量化并存入 ChromaDB，这可能需要一些时间...")
        # ChromaDB 有单次写入批次限制 (通常为 5461)，因此我们需要分批写入
        batch_size = 5000
        for i in range(0, len(new_docs), batch_size):
            batch = new_docs[i : i + batch_size]
            print(f"  正在分批写入第 {i} 到 {min(i + batch_size, len(new_docs))} 个片段...")
            retriever.store.add_documents(batch)
        print(f"✅ 成功！所有法律法规已经存入法规库 ({config.legal_db_dir})。")
    else:
        print("⚠️ 没有提取到有效文本。")

if __name__ == "__main__":
    main()
